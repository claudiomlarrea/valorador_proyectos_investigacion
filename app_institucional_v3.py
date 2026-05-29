import io
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

import pandas as pd
import streamlit as st
import yaml

# Lectura de archivos
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    from docx import Document
except Exception:
    Document = None


from valoracion_contenido import build_section_slices, score_project

_APP_DIR = Path(__file__).resolve().parent

APP_TITLE = "UCCuyo · Valorador de Proyectos de Investigación"
APP_VERSION = "v4.1 plantillas reales UCC + respaldo informe cátedra"

_CRITERIOS_V3_FALLBACK = {
    "Pertinencia y relevancia": {
        "peso": 10,
        "pistas": ["justificación", "relevancia", "problema", "fundamentación", "necesidad"]
    },
    "Claridad del problema y objetivos": {
        "peso": 10,
        "pistas": ["objetivo general", "objetivos específicos", "pregunta de investigación", "problema", "hipótesis"]
    },
    "Originalidad / aporte": {
        "peso": 8,
        "pistas": ["estado del arte", "marco teórico", "antecedentes", "novedad", "aporte", "vacancia"]
    },
    "Solidez metodológica": {
        "peso": 14,
        "pistas": ["metodología", "diseño", "enfoque", "técnicas", "análisis de datos", "método"]
    },
    "Calidad de datos / muestra": {
        "peso": 10,
        "pistas": ["muestra", "muestreo", "población", "instrumento", "datos", "recolección"]
    },
    "Factibilidad y cronograma": {
        "peso": 8,
        "pistas": ["cronograma", "plan de actividades", "factibilidad", "recursos", "viabilidad", "etapas"]
    },
    "Consideraciones éticas": {
        "peso": 6,
        "pistas": ["ética", "consentimiento", "confidencialidad", "comité de ética", "resguardo de datos"]
    },
    "Impacto esperado": {
        "peso": 8,
        "pistas": ["impacto", "resultados esperados", "beneficios", "relevancia social", "aportes"]
    },
    "Plan de difusión / transferencia": {
        "peso": 6,
        "pistas": ["difusión", "transferencia", "publicaciones", "divulgación", "congreso", "artículo"]
    },
    "Presupuesto y sostenibilidad": {
        "peso": 6,
        "pistas": ["presupuesto", "financiamiento", "costos", "recursos", "gastos", "sostenibilidad"]
    },
    "Alineación institucional y normativa": {
        "peso": 6,
        "pistas": ["institucional", "normativa", "lineamientos", "universidad", "facultad", "plan estratégico"]
    },
    "Bibliografía actualizada": {
        "peso": 8,
        "pistas": ["bibliografía", "referencias", "2021", "2022", "2023", "2024", "2025", "2026"]
    },
}


def _criterios_v3_validados_desde_yaml(data: Any) -> Optional[Dict]:
    """Si el YAML coincide con el respaldo (pesos y pistas), devuelve dict ordenado; si no, None."""
    if not isinstance(data, dict):
        return None
    if "criterios" in data and isinstance(data["criterios"], dict):
        data = data["criterios"]
    fb = _CRITERIOS_V3_FALLBACK
    if set(data.keys()) != set(fb.keys()):
        return None
    total = sum(int((data[k] or {}).get("peso", 0)) for k in data)
    if total != sum(int(v["peso"]) for v in fb.values()):
        return None
    for nombre, meta in fb.items():
        cm = data.get(nombre)
        if not isinstance(cm, dict):
            return None
        if int(cm.get("peso", -1)) != int(meta["peso"]):
            return None
        pistas_yaml = cm.get("pistas")
        if not isinstance(pistas_yaml, list):
            return None
        pistas_yaml_s = [str(x) for x in pistas_yaml]
        pistas_fb_s = [str(x) for x in meta["pistas"]]
        if pistas_yaml_s != pistas_fb_s:
            return None
    # Mismo contenido que el fallback → orden estable de la UI
    return {k: dict(data[k]) for k in fb.keys()}


def _load_criterios_v3() -> dict:
    path = _APP_DIR / "criterios_v3.yaml"
    if not path.is_file():
        return {k: dict(v) for k, v in _CRITERIOS_V3_FALLBACK.items()}
    try:
        raw = yaml.safe_load(path.read_text(encoding="utf-8"))
    except Exception:
        return {k: dict(v) for k, v in _CRITERIOS_V3_FALLBACK.items()}
    validado = _criterios_v3_validados_desde_yaml(raw)
    if validado is not None:
        return validado
    return {k: dict(v) for k, v in _CRITERIOS_V3_FALLBACK.items()}


CRITERIOS = _load_criterios_v3()


def categoria(p):
    if p >= 70:
        return "Aprobado"
    elif p >= 50:
        return "Aprobado con observaciones"
    elif p >= 30:
        return "Requiere reformulación"
    return "No aprobado"


def parse_pdf(file_bytes):
    if pdfplumber is None:
        return ""
    partes = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            partes.append(page.extract_text() or "")
    return "\n".join(partes)


def parse_docx(file_bytes):
    if DocxDocument is None:
        return ""
    doc = DocxDocument(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def make_excel(scores, porcentaje, resultado, nombre):
    filas = []
    for c, v in scores.items():
        filas.append({"Criterio": c, "Puntaje": v})

    df = pd.DataFrame(filas)

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Resultados")

        resumen = pd.DataFrame([{
            "Archivo": nombre,
            "Resultado": resultado,
            "Porcentaje": round(porcentaje, 2),
            "Fecha": datetime.now().strftime("%Y-%m-%d %H:%M")
        }])
        resumen.to_excel(writer, sheet_name="Resumen", index=False)

    return output.getvalue()


def make_word(scores, porcentaje, resultado, nombre):
    if Document is None:
        return b""

    doc = Document()
    doc.add_heading("Valoración de Proyecto de Investigación", 1)
    doc.add_paragraph(f"Archivo: {nombre}")
    doc.add_paragraph(f"Resultado: {resultado}")
    doc.add_paragraph(f"Cumplimiento: {round(porcentaje, 2)}%")

    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Criterio"
    hdr[1].text = "Puntaje"

    for c, v in scores.items():
        row = table.add_row().cells
        row[0].text = c
        row[1].text = str(v)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# ================= UI =================

st.set_page_config(page_title=APP_TITLE, layout="wide")
st.markdown("""
<style>

/* ================= FONDO ================= */
.stApp {
    background-color: #E6E6E6;
}

/* ================= TÍTULOS ================= */
.stApp h1, .stApp h2, .stApp h3, .stApp h4 {
    color: #064a3f;
}

/* ================= HEADER ================= */
.header-uccuyo {
    color: white !important;
}
.header-uccuyo * {
    color: white !important;
}

/* Ancho completo dentro del bloque principal (antes: max-width 900px inline centraba el banner) */
.header-uccuyo.banner-ucc-inst {
    width: 100% !important;
    max-width: 100% !important;
    box-sizing: border-box !important;
    margin-left: 0 !important;
    margin-right: 0 !important;
}

/* ================= TEXTO GENERAL ================= */
.stApp p,
.stApp label {
    color: #1a1a1a !important;
}

/* ================= UPLOADER ================= */
[data-testid="stFileUploader"] {
    background-color: white;
    border-radius: 10px;
    padding: 15px;
}

/* 🔥 BOTÓN UPLOAD (CORRECTO) */
[data-testid="stFileUploader"] button[kind="secondary"] {
    background-color: #064a3f !important;
    color: white !important;
    border-radius: 8px !important;
    border: none !important;
}

[data-testid="stFileUploader"] button[kind="secondary"]:hover {
    background-color: #0B6B5D !important;
}

/* ================= BOTONES GENERALES ================= */
.stButton button {
    background-color: #064a3f;
    color: white;
    border-radius: 8px;
    border: none;
    font-weight: 600;
}

.stButton button:hover {
    background-color: #0B6B5D;
}

/* ================= DOWNLOAD ================= */
[data-testid="stDownloadButton"] button {
    background-color: #064a3f !important;
    color: white !important;
    border-radius: 8px;
    border: none;
}

[data-testid="stDownloadButton"] button:hover {
    background-color: #0B6B5D !important;
}

/* 🔥 TEXTO BLANCO EN TODOS LOS BOTONES */
.stButton button *,
[data-testid="stDownloadButton"] button *,
[data-testid="stFileUploader"] button * {
    color: white !important;
}

/* ================= ALERTAS ================= */
[data-testid="stAlert"] {
    border-radius: 10px;
}

</style>
""", unsafe_allow_html=True)

st.markdown(
    """<div class="header-uccuyo banner-ucc-inst" style="background: linear-gradient(90deg, #0b5d4b, #177e6c); padding: 30px; border-radius: 15px; margin: 0 0 24px 0;">
<h1 style="margin:0;">Universidad Católica de Cuyo</h1>
<h2 style="margin-top:10px;">Secretaría de Investigación</h2>
<h3 style="margin-top:5px;">Consejo de Investigación</h3>
</div>""",
    unsafe_allow_html=True
)
        
st.title(APP_TITLE)
st.caption(APP_VERSION)

archivo = st.file_uploader("Subir proyecto (PDF o DOCX)", type=["pdf", "docx"])

if archivo is None:
    st.info("Esperando archivo…")
    st.stop()

raw = archivo.read()

texto = ""
if archivo.name.lower().endswith(".pdf"):
    texto = parse_pdf(raw)
else:
    texto = parse_docx(raw)

if texto.strip():
    st.success("Archivo cargado correctamente")
else:
    st.warning("Se cargó el archivo, pero no se extrajo texto visible.")

_scored = score_project(texto, CRITERIOS)
if len(_scored) == 2:
    criterion_results, section_slices = _scored
    _, doc_mode = build_section_slices(texto)
else:
    criterion_results, section_slices, doc_mode = _scored

st.info(
    f"**Modo detectado:** `{doc_mode}` — "
    "Proyecto ex ante (plantilla convocatoria) o informe final de cátedra. "
    "Si los apartados no coinciden, se usa texto propio de todo el documento (sin consignas)."
)
if doc_mode == "informe_catedra":
    st.warning(
        "Este archivo parece un **informe final de cátedra**, no un proyecto ex ante de convocatoria. "
        "La rúbrica del Anexo IV está pensada para la **plantilla de presentación de proyecto**. "
        "El puntaje automático es orientativo; conviene revisar con los sliders."
    )

with st.expander("Diagnóstico de apartados (contenido propio vs consignas)", expanded=False):
    if not section_slices:
        st.write("No se detectaron encabezados de apartado; se evalúa el texto como bloque único.")
    else:
        rows = []
        for key, sl in sorted(section_slices.items(), key=lambda x: x[0]):
            rows.append({
                "Apartado": key,
                "Palabras (propias)": sl.word_count_own,
                "% líneas consigna": round(sl.consigna_ratio * 100, 1),
            })
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

st.subheader("Evaluación")
st.caption(
    "v4.1: contenido propio por apartado (consignas no suman). "
    "Probar en local: `python3 -m streamlit run app_institucional_v3.py` → http://localhost:8501"
)

scores = {}
total_max = sum(meta["peso"] for meta in CRITERIOS.values())

cols = st.columns(2)
i = 0

for criterio, meta in CRITERIOS.items():
    with cols[i % 2]:
        peso = meta["peso"]
        cr = criterion_results[criterio]
        valor_inicial = cr.puntaje

        st.markdown(f"**{criterio}** (máx {peso})")
        st.caption(
            f"Ítems sustantivos: {cr.checks_ok}/{cr.checks_total} · "
            f"Palabras propias: {cr.word_own} (orient. {cr.word_target_min}–{cr.word_target_max})"
        )

        val = st.slider(
            f"Puntaje {criterio}",
            0,
            peso,
            valor_inicial,
            key=f"s_{i}"
        )

        with st.expander("Detalle automático"):
            for n in cr.notas:
                st.write(f"• {n}")
            if cr.evidencias:
                st.markdown("**Evidencia (contenido propio):**")
                for ev in cr.evidencias:
                    st.write(ev)

        scores[criterio] = val
        st.divider()

    i += 1

total = sum(scores.values())
porcentaje = (total / total_max) * 100
resultado = categoria(porcentaje)

st.markdown(f"## Resultado: **{resultado}**")
st.markdown(f"### Cumplimiento: **{round(porcentaje,2)}%**")

c1, c2 = st.columns(2)

with c1:
    st.download_button(
        "⬇️ Descargar Excel",
        make_excel(scores, porcentaje, resultado, archivo.name),
        "resultado.xlsx"
    )

with c2:
    st.download_button(
        "⬇️ Descargar Word",
        make_word(scores, porcentaje, resultado, archivo.name),
        "resultado.docx"
    )
