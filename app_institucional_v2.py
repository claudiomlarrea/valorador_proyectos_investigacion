
import io, hashlib
from datetime import datetime
import streamlit as st
import pandas as pd

# Parsing libs
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None

APP_TITLE = "UCCuyo · Valorador de Proyectos (v2 – auto 0 si no hay evidencia)"
APP_VERSION = "v2.0"

# --- Criterios y pesos por defecto (ex ante) ---
DEFAULT_CRITERIA = {
  "Pertinencia y relevancia": {"peso": 10, "pistas": ["justificación","relevancia","problema"]},
  "Claridad del problema y objetivos": {"peso": 10, "pistas": ["planteamiento del problema","objetivo general","objetivos específicos","pregunta de investigación"]},
  "Originalidad / aporte": {"peso": 8, "pistas": ["estado del arte","marco teórico","novedad"]},
  "Solidez metodológica": {"peso": 14, "pistas": ["metodología","diseño","enfoque","técnicas de análisis"]},
  "Calidad de datos / muestra": {"peso": 10, "pistas": ["datos","muestra","muestreo","instrumentos"]},
  "Factibilidad y cronograma": {"peso": 8, "pistas": ["cronograma","plan de actividades","recursos"]},
  "Consideraciones éticas": {"peso": 6, "pistas": ["ética","consentimiento","privacidad","comité de ética"]},
  "Impacto esperado": {"peso": 8, "pistas": ["resultados esperados","impacto","relevancia social"]},
  "Plan de difusión / transferencia": {"peso": 6, "pistas": ["plan de difusión","transferencia","publicaciones","congreso","artículo"]},
  "Presupuesto y sostenibilidad": {"peso": 6, "pistas": ["presupuesto","recursos","financiamiento","partida","costo","costos"]},
  "Alineación institucional y normativa": {"peso": 6, "pistas": ["institucional","lineamientos","normativa","política"]},
  "Bibliografía actualizada": {"peso": 8, "pistas": ["bibliografía","referencias","2021","2022","2023","2024","2025"]}
}

THRESHOLDS = {
    "Aprobado": (60, 1000),
    "Aprobado con observaciones": (50, 60),
    "No aprobado": (0, 50)
}

def categorize(porcentaje: float) -> str:
    for label, (lo, hi) in THRESHOLDS.items():
        if lo <= porcentaje < hi or (label == "Aprobado" and porcentaje == hi):
            return label
    return "No clasificado"

def parse_docx(file_bytes: bytes) -> str:
    if DocxDocument is None:
        return ""
    bio = io.BytesIO(file_bytes)
    doc = DocxDocument(bio)
    return "\n".join(p.text for p in doc.paragraphs)

def parse_pdf(file_bytes: bytes) -> str:
    if pdfplumber is None:
        return ""
    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for pg in pdf.pages:
            parts.append(pg.extract_text() or "")
    return "\n".join(parts)

def keyword_present(text_low: str, pistas):
    return any(p.lower() in text_low for p in pistas)

def score_ui(criteria_cfg, text: str):
    total_peso = sum(int(v.get("peso", 0)) for v in criteria_cfg.values())
    st.caption(f"Puntaje total posible: {total_peso} puntos (suma de pesos)")
    puntajes = {}
    cols = st.columns(2)
    i = 0
    text_low = text.lower()
    for crit, meta in criteria_cfg.items():
        with cols[i % 2]:
            peso = int(meta.get("peso", 0))
            st.markdown(f"**{crit}**  (peso {peso})")
            pistas = meta.get("pistas", [])
            hay = keyword_present(text_low, pistas)
            if not hay:
                st.info("No se detectó evidencia textual para este criterio → **Puntaje automático: 0**")
                val = 0
                obs = st.text_area("Observaciones", key=f"obs_{crit}", value="Sin evidencia textual en el documento.")
            else:
                val = st.slider("Puntaje asignado", 0, peso, int(round(peso*0.7)), key=f"score_{crit}")
                obs = st.text_area("Observaciones", key=f"obs_{crit}", placeholder="Notas, fortalezas, debilidades, recomendaciones…")
            puntajes[crit] = {"asignado": val, "peso": peso, "observaciones": obs}
            st.divider()
        i += 1
    obtenido = sum(v["asignado"] for v in puntajes.values())
    porcentaje = (obtenido / total_peso) * 100 if total_peso else 0.0
    return puntajes, obtenido, porcentaje, total_peso

def make_excel(criteria_cfg, puntajes, porcentaje, result, nombre_archivo):
    total_peso = sum(int(v.get("peso",0)) for v in criteria_cfg.values())
    rows = []
    for crit, meta in criteria_cfg.items():
        peso = int(meta.get("peso", 0))
        asignado = puntajes[crit]["asignado"]
        rows.append({
            "Criterio": crit,
            "Peso": peso,
            "Puntaje asignado": asignado,
            "Aporte (%)": round((asignado / total_peso)*100, 2) if total_peso>0 else 0,
            "Observaciones": puntajes[crit]["observaciones"]
        })
    df = pd.DataFrame(rows)
    with io.BytesIO() as output:
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="Resultados")
            resumen = pd.DataFrame([{
                "Archivo": nombre_archivo,
                "Resultado": result,
                "Porcentaje total": round(porcentaje,2),
                "Fecha": datetime.now().strftime("%Y-%m-%d %H:%M")
            }])
            resumen.to_excel(writer, index=False, sheet_name="Resumen")
        return output.getvalue()

def make_word(criteria_cfg, puntajes, porcentaje, result, nombre_archivo, extracto):
    if Document is None:
        return b""
    doc = Document()
    styles = doc.styles['Normal']
    styles.font.name = 'Times New Roman'
    styles.font.size = Pt(11)

    h = doc.add_paragraph("Universidad Católica de Cuyo – Secretaría de Investigación")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Valoración de Proyecto de Investigación", level=1)
    meta = doc.add_paragraph()
    meta.add_run("Archivo: ").bold = True
    meta.add_run(nombre_archivo + "   ")
    meta.add_run("Fecha: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))

    doc.add_paragraph(f"Dictamen: {result} — Cumplimiento: {round(porcentaje,2)}%")

    doc.add_heading("Resultados por criterio", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Criterio"
    hdr[1].text = "Peso"
    hdr[2].text = "Puntaje asignado"
    hdr[3].text = "Observaciones"

    for crit, meta in criteria_cfg.items():
        row = table.add_row().cells
        row[0].text = crit
        row[1].text = str(int(meta.get("peso",0)))
        row[2].text = str(puntajes[crit]["asignado"])
        row[3].text = puntajes[crit]["observaciones"] or ""

    fortalezas = [c for c in criteria_cfg if puntajes[c]["asignado"] >= int(criteria_cfg[c]["peso"]*0.75)]
    ausencias = [c for c in criteria_cfg if puntajes[c]["asignado"] == 0]

    doc.add_paragraph("")
    doc.add_heading("Síntesis", level=2)
    doc.add_paragraph("Fortalezas: " + (", ".join(fortalezas) if fortalezas else "No se destacan fortalezas específicas."))
    doc.add_paragraph("Ausencias/Aspectos a mejorar: " + (", ".join(ausencias) if ausencias else "Sin ausencias detectadas por palabras clave."))

    doc.add_paragraph("")
    doc.add_heading("Extracto de evidencia del documento", level=2)
    doc.add_paragraph(extracto[:2000])

    with io.BytesIO() as buffer:
        doc.save(buffer)
        return buffer.getvalue()

# ================== UI ==================
st.set_page_config(page_title=APP_TITLE, page_icon="🧮", layout="wide")
st.title(APP_TITLE)
st.caption("Cargá un proyecto (PDF/DOCX). Si un criterio no aparece en el texto, el sistema asigna **0 automático**. Exporta Excel y Word.")

uploaded = st.file_uploader("Proyecto (PDF o DOCX)", type=["pdf", "docx"])
if uploaded is None:
    st.info("Esperando archivo…")
    st.stop()

raw = uploaded.read()
sha = hashlib.sha256(raw).hexdigest()

if uploaded.name.lower().endswith(".pdf"):
    if pdfplumber is None:
        st.error("Falta dependencia: pdfplumber")
        st.stop()
    text = parse_pdf(raw)
else:
    if DocxDocument is None:
        st.error("Falta dependencia: python-docx")
        st.stop()
    text = parse_docx(raw)

criteria_cfg = DEFAULT_CRITERIA.copy()
puntajes, obtenido, porcentaje, total_peso = score_ui(criteria_cfg, text)
resultado = categorize(porcentaje)
st.markdown(f"### Resultado: **{resultado}** — Cumplimiento **{round(porcentaje,2)}%**")

c1, c2 = st.columns(2)
with c1:
    if st.button("⬇️ Exportar Excel"):
        xls = make_excel(criteria_cfg, puntajes, porcentaje, resultado, uploaded.name)
        st.download_button("Descargar resultados.xlsx", data=xls, file_name="valoracion_proyecto.xlsx")
with c2:
    if st.button("⬇️ Exportar Word"):
        docx_bytes = make_word(criteria_cfg, puntajes, porcentaje, resultado, uploaded.name, text[:4000])
        st.download_button("Descargar dictamen.docx", data=docx_bytes, file_name="dictamen_proyecto.docx")

st.divider()
st.caption("La detección usa palabras clave por criterio; si no aparecen, el puntaje se fija en 0. Editá listas de 'pistas' en el código si necesitás afinar por área.")
