import io
from datetime import datetime
import streamlit as st
import pandas as pd

# Librerías de lectura
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None

APP_TITLE = "UCCuyo · Valorador de Proyectos de Investigación"
APP_VERSION = "v2.1 – valoración flexible ajustada"

DEFAULT_CRITERIA = {
    "Pertinencia y relevancia": {
        "peso": 10,
        "pistas": ["justificación", "relevancia", "problema"]
    },
    "Claridad del problema y objetivos": {
        "peso": 10,
        "pistas": [
            "planteamiento del problema",
            "objetivo general",
            "objetivos específicos",
            "pregunta de investigación"
        ]
    },
    "Originalidad / aporte": {
        "peso": 8,
        "pistas": ["estado del arte", "marco teórico", "novedad", "aporte"]
    },
    "Solidez metodológica": {
        "peso": 14,
        "pistas": ["metodología", "diseño", "enfoque", "técnicas de análisis"]
    },
    "Calidad de datos / muestra": {
        "peso": 10,
        "pistas": ["datos", "muestra", "muestreo", "instrumentos"]
    },
    "Factibilidad y cronograma": {
        "peso": 8,
        "pistas": ["cronograma", "plan de actividades", "recursos", "viabilidad"]
    },
    "Consideraciones éticas": {
        "peso": 6,
        "pistas": ["ética", "consentimiento", "privacidad", "comité de ética", "confidencialidad"]
    },
    "Impacto esperado": {
        "peso": 8,
        "pistas": ["resultados esperados", "impacto", "relevancia social", "beneficios"]
    },
    "Plan de difusión / transferencia": {
        "peso": 6,
        "pistas": ["plan de difusión", "transferencia", "publicaciones", "artículo", "congreso", "divulgación"]
    },
    "Presupuesto y sostenibilidad": {
        "peso": 6,
        "pistas": ["presupuesto", "financiamiento", "recursos", "costos", "gastos", "sostenibilidad"]
    },
    "Alineación institucional y normativa": {
        "peso": 6,
        "pistas": ["institucional", "lineamientos", "normativa", "política", "universidad"]
    },
    "Bibliografía actualizada": {
        "peso": 8,
        "pistas": ["bibliografía", "referencias", "2021", "2022", "2023", "2024", "2025", "2026"]
    }
}

THRESHOLDS = {
    "Aprobado": (70, 1000),
    "Aprobado con observaciones": (50, 70),
    "Requiere reformulación": (30, 50),
    "No aprobado": (0, 30)
}


def categorize(p):
    for label, (lo, hi) in THRESHOLDS.items():
        if lo <= p < hi or (label == "Aprobado" and p == hi):
            return label
    return "No clasificado"


def parse_docx(file_bytes: bytes) -> str:
    if DocxDocument is None:
        return ""
    bio = io.BytesIO(file_bytes)
    doc = DocxDocument(bio)
    return "\n".join(p.text for p in doc.paragraphs)


def parse_pdf(file_bytes: bytes) -> str:
    if pdfplumber is None:
        return ""
    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for pg in pdf.pages:
            parts.append(pg.extract_text() or "")
    return "\n".join(parts)


def evidence_present(text: str, pistas: list[str]) -> bool:
    low = text.lower()
    for kw in pistas:
        if kw.lower() in low:
            return True
    return False


def extract_snippets(text: str, pistas: list[str], max_snippets: int = 2) -> list[str]:
    low = text.lower()
    snippets = []

    for kw in pistas:
        idx = low.find(kw.lower())
        if idx != -1:
            a = max(0, idx - 100)
            b = min(len(text), idx + 180)
            snippet = text[a:b].replace("\n", " ").strip()
            if snippet and snippet not in snippets:
                snippets.append(snippet)

        if len(snippets) >= max_snippets:
            break

    return snippets


def suggested_score(peso: int, hay_evidencia: bool) -> int:
    """
    Ajuste menos exigente:
    - Con evidencia: 85% del peso
    - Sin evidencia clara: 60% del peso
    """
    if hay_evidencia:
        return max(0, min(peso, int(round(peso * 0.85))))
    return max(0, min(peso, int(round(peso * 0.60))))


def score_ui(criteria_cfg: dict, text: str):
    total_peso = sum(int(v.get("peso", 0)) for v in criteria_cfg.values())
    st.caption(f"Puntaje total posible: {total_peso} puntos")

    puntajes = {}
    cols = st.columns(2)
    i = 0

    for crit, meta in criteria_cfg.items():
        with cols[i % 2]:
            peso = int(meta.get("peso", 0))
            pistas = meta.get("pistas", [])
            hay = evidence_present(text, pistas)
            snippets = extract_snippets(text, pistas)
            default_value = suggested_score(peso, hay)

            st.markdown(f"**{crit}**  (peso {peso})")

            if hay:
                st.success("Se detectó evidencia textual orientativa para este criterio.")
                obs_default = ""
            else:
                st.warning("No se detectó evidencia textual clara. Se recomienda revisión manual del evaluador.")
                obs_default = "La evidencia automática no fue clara; se recomienda revisión manual del evaluador."

            val = st.slider(
                "Puntaje asignado",
                min_value=0,
                max_value=peso,
                value=default_value,
                key=f"slider_{crit}"
            )

            obs = st.text_area(
                "Observaciones",
                key=f"obs_{crit}",
                value=obs_default,
                placeholder="Notas, fortalezas, debilidades, recomendaciones…"
            )

            if snippets:
                with st.expander("Evidencia sugerida", expanded=False):
                    for s in snippets:
                        st.code(s, language="markdown")

            puntajes[crit] = {
                "asignado": val,
                "peso": peso,
                "observaciones": obs
            }

            st.divider()
        i += 1

    obtenido = sum(v["asignado"] for v in puntajes.values())
    porcentaje = (obtenido / total_peso) * 100 if total_peso else 0.0
    return puntajes, obtenido, porcentaje, total_peso


def make_excel(criteria_cfg, puntajes, porcentaje, result, nombre_archivo):
    total_peso = sum(int(v.get("peso", 0)) for v in criteria_cfg.values())
    rows = []

    for crit, meta in criteria_cfg.items():
        peso = int(meta.get("peso", 0))
        asignado = puntajes[crit]["asignado"]

        rows.append({
            "Criterio": crit,
            "Peso": peso,
            "Puntaje asignado": asignado,
            "Aporte (%)": round((asignado / total_peso) * 100, 2) if total_peso > 0 else 0,
            "Observaciones": puntajes[crit]["observaciones"]
        })

    df = pd.DataFrame(rows)

    with io.BytesIO() as output:
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="Resultados")

            resumen = pd.DataFrame([{
                "Archivo": nombre_archivo,
                "Resultado": result,
                "Porcentaje total": round(porcentaje, 2),
                "Fecha": datetime.now().strftime("%Y-%m-%d %H:%M")
            }])
            resumen.to_excel(writer, index=False, sheet_name="Resumen")

        return output.getvalue()


def make_word(criteria_cfg, puntajes, porcentaje, result, nombre_archivo):
    if Document is None:
        return b""

    doc = Document()
    styles = doc.styles["Normal"]
    styles.font.name = "Times New Roman"
    styles.font.size = Pt(11)

    h = doc.add_paragraph("Universidad Católica de Cuyo – Secretaría de Investigación")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Valoración de Proyecto de Investigación", level=1)

    meta = doc.add_paragraph()
    meta.add_run("Archivo: ").bold = True
    meta.add_run(nombre_archivo + "   ")
    meta.add_run("Fecha: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))

    doc.add_paragraph(f"Dictamen: {result} — Cumplimiento: {round(porcentaje, 2)}%")

    doc.add_heading("Resultados por criterio", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Criterio"
    hdr[1].text = "Peso"
    hdr[2].text = "Puntaje asignado"
    hdr[3].text = "Observaciones"

    for crit, meta in criteria_cfg.items():
        row = table.add_row().cells
        row[0].text = crit
        row[1].text = str(int(meta.get("peso", 0)))
        row[2].text = str(puntajes[crit]["asignado"])
        row[3].text = puntajes[crit]["observaciones"] or ""

    fortalezas = [
        c for c in criteria_cfg
        if puntajes[c]["asignado"] >= int(criteria_cfg[c]["peso"] * 0.75)
    ]
    debiles = [
        c for c in criteria_cfg
        if puntajes[c]["asignado"] <= int(criteria_cfg[c]["peso"] * 0.40)
    ]

    doc.add_paragraph("")
    doc.add_heading("Síntesis", level=2)
    doc.add_paragraph(
        "Fortalezas: " +
        (", ".join(fortalezas) if fortalezas else "No se destacan fortalezas específicas.")
    )
    doc.add_paragraph(
        "Aspectos a revisar: " +
        (", ".join(debiles) if debiles else "No se identifican debilidades críticas.")
    )

    with io.BytesIO() as buffer:
        doc.save(buffer)
        return buffer.getvalue()


# ================== UI ==================
st.set_page_config(page_title=APP_TITLE, page_icon="🧮", layout="wide")
st.title(APP_TITLE)
st.caption(f"{APP_VERSION}. La detección automática de evidencia es orientativa y el evaluador puede ajustar manualmente el puntaje.")

uploaded = st.file_uploader("Proyecto (PDF o DOCX)", type=["pdf", "docx"])

if uploaded is None:
    st.info("Esperando archivo…")
    st.stop()

raw = uploaded.read()

if uploaded.name.lower().endswith(".pdf"):
    if pdfplumber is None:
        st.error("Falta dependencia: pdfplumber")
        st.stop()
    text = parse_pdf(raw)
else:
    if DocxDocument is None:
        st.error("Falta dependencia: python-docx")
        st.stop()
    text = parse_docx(raw)

criteria_cfg = DEFAULT_CRITERIA.copy()

puntajes, obtenido, porcentaje, total_peso = score_ui(criteria_cfg, text)
resultado = categorize(porcentaje)

st.markdown(f"### Resultado: **{resultado}** — Cumplimiento **{round(porcentaje, 2)}%**")

c1, c2 = st.columns(2)

with c1:
    xls = make_excel(criteria_cfg, puntajes, porcentaje, resultado, uploaded.name)
    st.download_button(
        "⬇️ Descargar resultados.xlsx",
        data=xls,
        file_name="valoracion_proyecto.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

with c2:
    docx_bytes = make_word(criteria_cfg, puntajes, porcentaje, resultado, uploaded.name)
    st.download_button(
        "⬇️ Descargar dictamen.docx",
        data=docx_bytes,
        file_name="dictamen_proyecto.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.divider()
st.caption("La detección automática de evidencia es orientativa. El dictamen final debe apoyarse en el juicio académico del evaluador.")
