# app_institucional_v3.py
import io, re, yaml, pdfplumber, os
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from datetime import datetime

# =========================
# Configuración general
# =========================
st.set_page_config(
    page_title="UCCuyo · Valorador de Proyectos de Investigación",
    page_icon="🧪",
    layout="wide"
)

# =========================
# Carga de rúbrica (tolerante a nombres)
# =========================
@st.cache_resource
def load_rubric():
    posibles = [
        "rubric_proyecto.yaml",
        "rubric_proyectos.yaml",
        "rubric_project.yaml",
        "rubric_projects.yaml",
        "rubric_institucional.yaml",
        "rubric_config.yaml"
    ]
    ultimo_error = None
    for fname in posibles:
        if os.path.exists(fname):
            try:
                with open(fname, "r", encoding="utf-8") as f:
                    data = yaml.safe_load(f)
                    return data, fname
            except Exception as e:
                ultimo_error = e
                continue
    if ultimo_error:
        st.warning(f"No pude cargar ninguna rúbrica. Último error: {ultimo_error}")
    # Rúbrica mínima por defecto si no hay archivo
    default = {
        "scale": {"max": 4},
        "weights": {
            "identificacion": 5,
            "equipo": 8,
            "justificacion": 12,
            "objetivos": 12,
            "marco_teorico": 10,
            "metodologia": 15,
            "plan_actividades": 10,
            "cronograma": 6,
            "presupuesto": 8,
            "resultados_esperados": 8,
            "difusion_transferencia": 4,
            "viabilidad_etica": 6,
            "calidad_formal": 6
        },
        "thresholds": {"aprobado": 60, "aprobado_obs": 50},
        "keywords": {}
    }
    return default, None

RUBRIC, RUBRIC_FILE = load_rubric()

# Si la rúbrica define pesos, construir la lista de criterios desde weights (para no desalinear claves)
def _formatear_nombre(key: str) -> str:
    return key.replace("_", " ").replace("-", " ").strip().capitalize()

if "weights" in RUBRIC and isinstance(RUBRIC["weights"], dict) and len(RUBRIC["weights"]) > 0:
    CRITERIA = [(k, _formatear_nombre(k)) for k in RUBRIC["weights"].keys()]
else:
    # fallback estable (si se usa default)
    CRITERIA = [
        ("identificacion", "Identificación general del proyecto"),
        ("equipo", "Equipo de trabajo"),
        ("justificacion", "Fundamentación / justificación"),
        ("objetivos", "Objetivos"),
        ("marco_teorico", "Marco teórico"),
        ("metodologia", "Metodología"),
        ("plan_actividades", "Plan de actividades"),
        ("cronograma", "Cronograma"),
        ("presupuesto", "Presupuesto"),
        ("resultados_esperados", "Resultados esperados / impacto"),
        ("difusion_transferencia", "Difusión y transferencia"),
        ("viabilidad_etica", "Viabilidad y aspectos éticos"),
        ("calidad_formal", "Calidad formal"),
    ]

MAX_SCALE = RUBRIC.get("scale", {}).get("max", 4)

# =========================
# Extracción de texto
# =========================
def extract_text_from_docx(file_bytes: bytes) -> str:
    buffer = io.BytesIO(file_bytes)
    doc = Document(buffer)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_pdf(file_bytes: bytes) -> str:
    buffer = io.BytesIO(file_bytes)
    text_parts = []
    with pdfplumber.open(buffer) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)

# =========================
# Scoring simple (heurístico por keywords)
# =========================
def naive_auto_score(text: str, key: str) -> int:
    words = RUBRIC.get("keywords", {}).get(key, [])
    if not words:
        return 0
    lower = text.lower()
    hits = sum(1 for w in words if w and w.lower() in lower)
    ratio = hits / max(len(words), 1)
    if ratio == 0:
        return 0
    elif ratio < 0.25:
        return 1
    elif ratio < 0.5:
        return 2
    elif ratio < 0.75:
        return 3
    else:
        return 4

def weighted_total(scores: dict) -> float:
    weights = RUBRIC.get("weights", {})
    total = 0.0
    for k, v in scores.items():
        w = weights.get(k, 0)
        total += (v / MAX_SCALE) * w
    return round(total, 2)

def decision(final_pct: float) -> str:
    th = RUBRIC.get("thresholds", {"aprobado": 60, "aprobado_obs": 50})
    if final_pct >= th.get("aprobado", 60):
        return "APROBADO"
    elif final_pct >= th.get("aprobado_obs", 50):
        return "APROBADO CON OBSERVACIONES"
    else:
        return "NO APROBADO"

# =========================
# Helpers Word (anti-truncado)
# =========================
def _split_long_paragraphs(text: str, max_len: int = 2000):
    """
    Divide un párrafo largo en trozos <= max_len para evitar el límite de 32.767
    caracteres por párrafo de Word. Intenta cortar en espacios para no partir palabras.
    """
    text = text.strip()
    if not text:
        return []
    chunks, i, n = [], 0, len(text)
    while i < n:
        j = min(i + max_len, n)
        k = text.rfind(" ", i, j)
        if k == -1 or k <= i + int(max_len*0.6):
            k = j
        chunk = text[i:k].strip()
        if chunk:
            chunks.append(chunk)
        i = k
    return chunks

def _add_full_text_as_paragraphs(doc: Document, text: str, max_len: int = 2000) -> None:
    """
    Inserta texto en párrafos limpios. Usa doble salto como separador preferente; si no hay,
    acepta salto simple. Cualquier párrafo que supere max_len se trocea en partes seguras.
    """
    if not text:
        return
    blocks = re.split(r"\n{2,}", text.strip()) if "\n\n" in text else text.split("\n")
    for block in blocks:
        block = " ".join([ln.strip() for ln in block.splitlines() if ln.strip()])
        if not block:
            doc.add_paragraph("")
            continue
        for chunk in _split_long_paragraphs(block, max_len=max_len):
            p = doc.add_paragraph(chunk)
            p.paragraph_format.space_after = Pt(6)

def _recortar_evidencia_proyecto(raw_text: str) -> str:
    """
    Devuelve el bloque de 'evidencia' a pegar en el Word.
    - Intenta comenzar desde encabezados típicos: 'Resumen del trabajo de investigación',
      'Resumen', 'DESCRIPCIÓN DEL PLAN DE TRABAJO FINAL', 'CARÁTULA'.
    - Si no encuentra, devuelve TODO el texto.
    No impone un fin: el troceo por párrafos evita el corte de Word.
    """
    if not raw_text:
        return raw_text
    inicios = [
        "Resumen del trabajo de investigación",
        "RESUMEN DEL TRABAJO DE INVESTIGACIÓN",
        "Resumen",
        "RESUMEN",
        "DESCRIPCION DEL PLAN DE TRABAJO FINAL",
        "DESCRIPCIÓN DEL PLAN DE TRABAJO FINAL",
        "CARATULA",
        "CARÁTULA",
        "Extracto de evidencia del documento"
    ]
    lower = raw_text.lower()
    start_pos = -1
    for patt in inicios:
        pos = lower.find(patt.lower())
        if pos != -1 and (start_pos == -1 or pos < start_pos):
            start_pos = pos
    return (raw_text[start_pos:].strip() if start_pos != -1 else raw_text.strip())

# =========================
# Exportaciones
# =========================
def make_excel(scores: dict, final_pct: float, label: str) -> bytes:
    weights = RUBRIC.get("weights", {})
    df = pd.DataFrame([{
        "Criterio": name,
        "Clave": key,
        "Puntaje (0-4)": scores.get(key, 0),
        "Peso (%)": weights.get(key, 0),
        "Aporte (%)": round((scores.get(key,0)/MAX_SCALE)*weights.get(key,0), 2)
    } for key, name in CRITERIA])
    df_total = pd.DataFrame([{"Total (%)": final_pct, "Dictamen": label}])
    with io.BytesIO() as output:
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="Resultados")
            df_total.to_excel(writer, index=False, sheet_name="Resumen")
        return output.getvalue()

def make_word(scores: dict, final_pct: float, label: str, raw_text: str) -> bytes:
    weights = RUBRIC.get("weights", {})
    doc = Document()

    # Estilo base
    styles = doc.styles['Normal']
    styles.font.name = 'Times New Roman'
    styles.font.size = Pt(11)

    # Márgenes amplios (más área útil)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Encabezado
    doc.add_heading('UCCuyo – Valoración de Proyecto de Investigación', level=1)
    today = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc.add_paragraph(f"Fecha: {today}")
    doc.add_paragraph(f"Dictamen: {label}  —  Cumplimiento: {final_pct}%")

    doc.add_paragraph("")
    doc.add_heading('Resultados por criterio', level=2)
    for key, name in CRITERIA:
        s = scores.get(key, 0)
        w = weights.get(key, 0)
        aporte = round((s/MAX_SCALE)*w, 2)
        p = doc.add_paragraph()
        run_title = p.add_run(f"{name} ")
        run_title.bold = True
        p.add_run(f"(Puntaje: {s}/{MAX_SCALE} · Peso: {w}% · Aporte: {aporte}%)")

    doc.add_paragraph("")
    doc.add_heading('Interpretación', level=2)
    fortalezas = [name for key, name in CRITERIA if scores.get(key,0) >= 3]
    mejoras = [name for key, name in CRITERIA if scores.get(key,0) <= 1]
    doc.add_paragraph("Fortalezas: " + (", ".join(fortalezas) if fortalezas else "no se identifican fortalezas destacadas."))
    doc.add_paragraph("Aspectos a mejorar: " + (", ".join(mejoras) if mejoras else "no se identifican aspectos críticos."))

    doc.add_paragraph("")
    doc.add_heading('Evidencia analizada (texto completo)', level=2)

    # Evidencia desde encabezados típicos (si existen) + troceo anti-truncado
    evidencia = _recortar_evidencia_proyecto(raw_text)
    _add_full_text_as_paragraphs(doc, evidencia, max_len=2000)

    with io.BytesIO() as buffer:
        doc.save(buffer)
        return buffer.getvalue()

# =========================
# Interfaz
# =========================
st.markdown("## 🧪 Valorador de Proyectos de Investigación")
st.write("Subí un **PDF o DOCX** del proyecto. La app extrae el texto, propone un puntaje automático por criterios configurables y te permite **ajustarlos manualmente** antes de exportar resultados.")

uploaded = st.file_uploader("Cargar archivo (PDF o DOCX)", type=["pdf", "docx"])

raw_text = ""
if uploaded is not None:
    data = uploaded.read()
    if uploaded.name.lower().endswith(".docx"):
        raw_text = extract_text_from_docx(data)
    else:
        raw_text = extract_text_from_pdf(data)

    with st.expander("📄 Texto extraído (vista previa)"):
        st.text_area("Contenido (vista previa)", raw_text[:6000], height=280)

    st.divider()
    st.subheader("Evaluación automática + ajuste manual")

    cols = st.columns(3)
    auto_scores = {}
    for idx, (key, name) in enumerate(CRITERIA):
        if idx % 3 == 0:
            cols = st.columns(3)
        col = cols[idx % 3]
        with col:
            auto = naive_auto_score(raw_text, key)
            auto_scores[key] = int(auto)

    st.write("**Sugerencia automática (0–4)**:", auto_scores)

    st.markdown("### Ajustar puntajes (0–4)")
    scores = {}
    for key, name in CRITERIA:
        scores[key] = st.slider(name, min_value=0, max_value=int(MAX_SCALE), value=int(auto_scores.get(key,0)))

    final_pct = weighted_total(scores)
    label = decision(final_pct)
    st.markdown(f"### Resultado: **{label}** — Cumplimiento **{final_pct}%**")

    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("⬇️ Exportar Excel"):
            xls = make_excel(scores, final_pct, label)
            st.download_button("Descargar resultados.xlsx", data=xls, file_name="valoracion_proyecto.xlsx")
    with c2:
        if st.button("⬇️ Exportar Word"):
            docx_bytes = make_word(scores, final_pct, label, raw_text)
            st.download_button("Descargar dictamen.docx", data=docx_bytes, file_name="dictamen_proyecto.docx")
    with c3:
        if RUBRIC_FILE and os.path.exists(RUBRIC_FILE):
            st.download_button("Descargar configuración (YAML)", data=open(RUBRIC_FILE,"rb").read(), file_name=RUBRIC_FILE)
        else:
            st.download_button("Descargar configuración (YAML por defecto)", data=yaml.safe_dump(RUBRIC).encode("utf-8"), file_name="rubric_proyecto_generada.yaml")
else:
    st.info("Esperando archivo...")
